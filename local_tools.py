# %%
import os
import sys
import docx
import webbrowser
import pandas as pd

from pprint import pprint

from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed

import PyPDF2

import logging
logging.propagate = False
logging.getLogger().setLevel(logging.ERROR)

EXTS = ['.docx', '.pdf']

# %%


def read_pdf(pdf):
    # Read contents of [pdf]
    praser = PDFParser(pdf)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器 与文档对象
    praser.set_document(doc)
    doc.set_parser(praser)

    # 提供初始化密码
    # 如果没有密码 就创建一个空的字符串
    doc.initialize()

    lines = []
    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        print('!!! not extractable.')
        return '[Not extractable]'
    else:
        # 创建PDf 资源管理器 来管理共享资源
        rsrcmgr = PDFResourceManager()
        # 创建一个PDF设备对象
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        # 循环遍历列表，每次处理一个page的内容
        for page in doc.get_pages():  # doc.get_pages() 获取page列表
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等 想要获取文本就获得对象的text属性，
            for x in layout:
                if (isinstance(x, LTTextBoxHorizontal)):
                    lines.append(x.get_text())

    # Return
    return '\n'.join(lines)


# %%
def get_text(path):
    # Get text from path

    # Get text of .docx file
    if path.endswith('.docx'):
        texts = []

        doc = docx.Document(path)
        all_paras = doc.paragraphs
        for para in all_paras:
            texts.append(para.text)

        return '\n'.join(texts)

    # Get text of .pdf file
    if path.endswith('.pdf'):
        texts = []

        with open(path, 'rb') as f:
            return read_pdf(f)
            pdf_reader = PyPDF2.PdfFileReader(f)
            for j in range(pdf_reader.numPages):
                page = pdf_reader.getPage(j)
                texts.append(page.extractText())

        return '\n'.join(texts)

    # Get nothing
    return ''


# %%
class Recorder():
    def __init__(self, words):
        self.suspects = pd.DataFrame(columns=['Name',
                                              'Offence',
                                              'Surrounding',
                                              'Folder'])

        self.words = words

    def append(self, dct):
        series = pd.Series(dct)
        self.suspects = self.suspects.append(series,
                                             ignore_index=True)
        pass

    def display(self):
        print('--------------------')
        pprint(self.suspects)

    def save(self, name='suspects', browser=False):
        self.suspects.to_json(f'{name}.json')
        self.suspects.to_html(f'{name}.html')
        if browser:
            webbrowser.open(f'{name}.html')


# %%

class FileManager():
    """FileManager

    Find all .docx files in root recursively.
    """

    def __init__(self, root, exts=EXTS):
        self.root = root
        self.exts = exts
        self.paths = []

    def append(self, path):
        # Append new [path] to paths
        # path: full path of the .docx file

        # Path should be a file
        assert(os.path.isfile(path))

        # Ignore temporal file
        if os.path.basename(path).startswith('~'):
            return

        # Record legal .docx file
        if any([path.endswith(ext) for ext in self.exts]):
            self.paths.append(path)

    def walker(self, root=None):
        # Walk through [root] recursively

        # Start with self.root
        if root is None:
            root = self.root

        # Use try wrapper to deal with PermissionError
        try:
            for node in os.listdir(root):
                # If meet a file, try to append it
                if os.path.isfile(os.path.join(root, node)):
                    self.append(os.path.join(root, node))

                # if meet a folder, dig in
                if os.path.isdir(os.path.join(root, node)):
                    self.walker(root=os.path.join(root, node))
        except PermissionError:
            pass

    def display(self):
        # Display recorded paths
        print('-' * 80)
        for j, path in enumerate(self.paths):
            print(j, path)
        print('')


# %%
if __name__ == '__main__':
    manager = FileManager()
    manager.walker()
    manager.display()

# %%
